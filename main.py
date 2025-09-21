from fastapi import FastAPI, UploadFile, File
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import tempfile, base64, os, re, requests, logging
from fastapi.middleware.cors import CORSMiddleware

from utills.blobservice import azure_blob_manager  # your blob manager

load_dotenv()
MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
if not MISTRAL_API_KEY:
    raise RuntimeError("Set MISTRAL_API_KEY in your .env")

MISTRAL_OCR_URL = "https://api.mistral.ai/v1/ocr"

app = FastAPI()

origins = [
    "https://dock-maker-ui.netlify.app",  # your frontend
    "http://nirmaldocai.com",              # local dev
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,          # or ["*"] for all origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
logging.basicConfig(level=logging.INFO)


def parse_markdown_tables(md: str):
    """Parse markdown-style tables into lists of rows and return cleaned markdown."""
    lines = md.splitlines()
    tables, out_lines = [], []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and line.count("|") >= 2:
            block = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].count("|") >= 2:
                block.append(lines[i])
                i += 1
            rows = []
            for r in block:
                row_text = r.strip()
                if row_text.startswith("|"): row_text = row_text[1:]
                if row_text.endswith("|"): row_text = row_text[:-1]
                cells = [cell.strip() for cell in row_text.split("|")]
                if all(re.fullmatch(r"-{1,}", c) or re.fullmatch(r":?-+:?", c) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                tables.append(rows)
            continue
        else:
            out_lines.append(line)
            i += 1
    return tables, "\n".join(out_lines)


def markdown_to_plaintext(md: str) -> str:
    """Convert markdown to plain text."""
    if not md:
        return ""
    md = re.sub(r"!\[.*?\]\(.*?\)", "", md)
    md = re.sub(r"^#{1,6}\s*", "", md, flags=re.MULTILINE)
    md = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", md)
    md = re.sub(r"<br\s*/?>", "\n", md, flags=re.IGNORECASE)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


@app.post("/convert/")
async def convert_pdf(file: UploadFile = File(...)):
    # Save uploaded PDF temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        pdf_path = tmp_pdf.name
        content = await file.read()
        tmp_pdf.write(content)

    output_path = pdf_path.replace(".pdf", ".docx")

    logging.info("Scanned PDF -> calling Mistral OCR")
    base64_pdf = base64.b64encode(content).decode("utf-8")
    payload = {
        "model": "mistral-ocr-latest",
        "document": {
            "type": "document_url",
            "document_url": f"data:application/pdf;base64,{base64_pdf}"
        },
        "include_image_base64": True
    }

    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }

    try:
        resp = requests.post(MISTRAL_OCR_URL, headers=headers, json=payload, timeout=180)
        resp.raise_for_status()
    except Exception as e:
        logging.exception("Mistral OCR request failed: %s", e)
        raise

    resp_json = resp.json()
    pages = resp_json.get("pages", [])
    doc = Document()

    for page in pages:
        md = page.get("markdown", "") or ""
        tables, md_without_tables = parse_markdown_tables(md)

        # Extract title: first heading
        title = None
        for line in md.splitlines():
            h = re.match(r"^\s{0,3}#{1,6}\s+(.*\S)", line)
            if h:
                title = h.group(1).strip()
                break

        if title:
            h = doc.add_heading(title, level=1)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add plain paragraphs
        plain_text = markdown_to_plaintext(md_without_tables)
        if plain_text:
            for para in [p.strip() for p in plain_text.split("\n\n") if p.strip()]:
                p = doc.add_paragraph(para)
                for run in p.runs:
                    run.font.size = Pt(10)

        # Add tables
        for table in tables:
            max_cols = max(len(r) for r in table)
            rows = len(table)
            if rows == 0: continue
            word_table = doc.add_table(rows=rows, cols=max_cols)
            word_table.style = "Table Grid"
            for r_idx, row in enumerate(table):
                for c_idx in range(max_cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    cell = word_table.cell(r_idx, c_idx)
                    cell.text = val
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
            doc.add_paragraph("")

        doc.add_page_break()

    doc.save(output_path)
    logging.info("Saved docx: %s", output_path)

    # Upload to Azure Blob Storage
    blob_name = os.path.basename(output_path)
    with open(output_path, "rb") as f:
        blob_url = azure_blob_manager.upload_file(blob_name=blob_name, file=f)

    logging.info("Uploaded DOCX to Blob Storage: %s", blob_url)

    return {"blob_url": blob_url}
